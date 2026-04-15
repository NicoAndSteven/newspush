"""
直接生成 Word 文档模块
不经过 Markdown 转换，直接从分析结果生成格式化的 Word 文档
"""
import requests
import io
from typing import List, Dict, Optional, TYPE_CHECKING
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # 定义一个占位符类型以避免类型注解错误
    if TYPE_CHECKING:
        from docx import Document
    else:
        Document = object


class DirectWordGenerator:
    """直接生成 Word 文档"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装")
    
    def generate_document(
        self,
        title: str,
        summary: str,
        core_facts: Dict,
        key_points: List[str],
        background: str,
        impact_analysis: str,
        unique_angle: str,
        controversial_aspects: List[str],
        expert_opinion: str,
        future_outlook: str,
        images: List[str],
        output_path: str
    ) -> str:
        """
        直接生成 Word 文档
        
        Args:
            title: 文档标题
            summary: 导语/摘要
            core_facts: 核心事实字典
            key_points: 核心要点列表
            background: 背景分析
            impact_analysis: 影响分析
            unique_angle: 独特视角
            controversial_aspects: 争议焦点
            expert_opinion: 专家点评
            future_outlook: 未来展望
            images: 图片 URL 列表
            output_path: 输出路径
        
        Returns:
            生成的文件路径
        """
        doc = Document()
        
        # 设置页面边距
        self._set_page_margins(doc)
        
        # 设置默认字体
        self._set_default_font(doc)
        
        # 1. 标题
        self._add_title(doc, title)
        
        # 2. 封面图
        if images:
            self._add_image(doc, images[0], width_inches=5.5)
        
        # 3. 导语
        if summary:
            self._add_lead_paragraph(doc, summary)
        
        # 4. 核心事实
        if core_facts:
            self._add_core_facts(doc, core_facts)
        
        # 5. 核心要点
        if key_points:
            self._add_key_points(doc, key_points)
        
        # 6. 第二张图
        if len(images) > 1:
            self._add_image(doc, images[1], width_inches=5.0)
        
        # 7. 背景与影响
        self._add_section(doc, "背景与影响", [
            ("背景", background),
            ("影响分析", impact_analysis)
        ])
        
        # 8. 深度分析
        self._add_deep_analysis(doc, unique_angle, controversial_aspects, expert_opinion)
        
        # 9. 第三张图
        if len(images) > 2:
            self._add_image(doc, images[2], width_inches=5.0)
        
        # 10. 未来展望
        if future_outlook:
            self._add_section(doc, "展望", [("未来趋势", future_outlook)])
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def _set_page_margins(self, doc: Document):
        """设置页面边距"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
    
    def _set_default_font(self, doc: Document):
        """设置默认字体"""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.space_after = Pt(6)
    
    def _add_title(self, doc: Document, title: str):
        """添加标题"""
        # 主标题
        heading = doc.add_heading(level=0)
        run = heading.add_run(title)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        doc.add_paragraph()
    
    def _add_lead_paragraph(self, doc: Document, text: str):
        """添加导语段落（加粗显示）"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(14)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进
        
        doc.add_paragraph()  # 空行
    
    def _add_core_facts(self, doc: Document, facts: Dict):
        """添加核心事实"""
        # 小标题
        self._add_subheading(doc, "核心事实")
        
        # 事实列表
        if facts.get('event_date'):
            self._add_fact_item(doc, "时间", facts['event_date'])
        if facts.get('location'):
            self._add_fact_item(doc, "地点", facts['location'])
        if facts.get('key_figures'):
            figures = facts['key_figures']
            if isinstance(figures, list):
                names = ', '.join([f.get('name', '') for f in figures if f.get('name')])
                if names:
                    self._add_fact_item(doc, "关键人物", names)
        if facts.get('main_event'):
            self._add_fact_item(doc, "事件", facts['main_event'])
        
        doc.add_paragraph()  # 空行
    
    def _add_fact_item(self, doc: Document, label: str, value: str):
        """添加事实项"""
        p = doc.add_paragraph()
        
        # 标签（粗体）
        run1 = p.add_run(f"{label}：")
        run1.font.name = 'Microsoft YaHei'
        run1.font.bold = True
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 值
        run2 = p.add_run(value)
        run2.font.name = 'Microsoft YaHei'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _add_key_points(self, doc: Document, points: List[str]):
        """添加核心要点"""
        self._add_subheading(doc, "核心要点")
        
        for i, point in enumerate(points, 1):
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(point)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        doc.add_paragraph()
    
    def _add_section(self, doc: Document, title: str, contents: List[tuple]):
        """添加章节"""
        self._add_subheading(doc, title)
        
        for label, content in contents:
            if content:
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                p.paragraph_format.first_line_indent = Cm(0.74)
        
        doc.add_paragraph()
    
    def _add_deep_analysis(self, doc: Document, unique_angle: str, controversial: List[str], opinion: str):
        """添加深度分析"""
        self._add_subheading(doc, "深度分析")
        
        # 独特视角
        if unique_angle:
            p = doc.add_paragraph()
            run = p.add_run(unique_angle)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            p.paragraph_format.first_line_indent = Cm(0.74)
        
        # 争议焦点
        if controversial:
            p = doc.add_paragraph()
            run = p.add_run("多方视角")
            run.font.name = 'Microsoft YaHei'
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            for aspect in controversial:
                bullet = doc.add_paragraph(style='List Bullet')
                run = bullet.add_run(aspect)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 专家点评
        if opinion:
            paragraphs = opinion.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(para.strip())
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    p.paragraph_format.first_line_indent = Cm(0.74)
        
        doc.add_paragraph()
    
    def _add_subheading(self, doc: Document, text: str):
        """添加小标题"""
        heading = doc.add_heading(level=2)
        run = heading.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _add_image(self, doc: Document, image_url: str, width_inches: float = 5.0):
        """添加图片"""
        try:
            if image_url.startswith('http'):
                response = requests.get(image_url, timeout=15)
                if response.status_code == 200:
                    image_stream = io.BytesIO(response.content)
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(image_stream, width=Inches(width_inches))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
            elif Path(image_url).exists():
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(image_url, width=Inches(width_inches))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        except Exception as e:
            print(f"    [警告] 插入图片失败: {e}")


def generate_word_directly(
    title: str,
    summary: str,
    core_facts: Dict,
    key_points: List[str],
    background: str,
    impact_analysis: str,
    unique_angle: str,
    controversial_aspects: List[str],
    expert_opinion: str,
    future_outlook: str,
    images: List[str],
    output_path: str
) -> str:
    """便捷函数：直接生成 Word 文档"""
    generator = DirectWordGenerator()
    return generator.generate_document(
        title=title,
        summary=summary,
        core_facts=core_facts,
        key_points=key_points,
        background=background,
        impact_analysis=impact_analysis,
        unique_angle=unique_angle,
        controversial_aspects=controversial_aspects,
        expert_opinion=expert_opinion,
        future_outlook=future_outlook,
        images=images,
        output_path=output_path
    )


if __name__ == "__main__":
    # 测试
    output = generate_word_directly(
        title="特朗普要求伊朗结束核计划，美国实施海上封锁",
        summary="美国总统特朗普宣布对伊朗实施海上封锁，要求其终止核计划。",
        core_facts={
            "event_date": "2026-04-13",
            "location": "波斯湾",
            "key_figures": [{"name": "特朗普", "title": "总统"}],
            "main_event": "美国对伊朗实施海上封锁"
        },
        key_points=[
            "美国宣布对伊朗实施全面海上封锁",
            "伊朗核计划成为焦点",
            "国际社会密切关注事态发展"
        ],
        background="这是自2018年美国退出伊朗核协议以来最严重的升级...",
        impact_analysis="此举可能引发地区局势进一步紧张...",
        unique_angle="从地缘政治角度看，这是美国在中东地区的重大战略调整...",
        controversial_aspects=[
            "封锁是否构成战争行为",
            "伊朗可能采取的反制措施"
        ],
        expert_opinion="专家认为，此次封锁将对全球能源市场产生深远影响...",
        future_outlook="未来局势发展取决于多方因素...",
        images=[],
        output_path="test_direct_word.docx"
    )
    print(f"Word 文档已生成: {output}")
